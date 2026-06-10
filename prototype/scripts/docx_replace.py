#!/usr/bin/env python3
"""docx_replace.py — 原本 docx をコピーし、合意した変更(変更前→変更後)だけを差し替える CLI。

資料レビュー(document-review)の改訂版生成で、**元の体裁・画像・チャートを保ったまま**
テキストだけを直す「原本コピー編集」方式を担うヘルパー(→ C2)。全面再構築(md_to_docx)が
画像・複雑な書式を失うのに対し、こちらは原本をベースに差し替えるので体裁が保たれる。

変更は JSON で渡す(段落・表セルを走査し、各 find を replace に置換):
  [{"find": "変更前の文字列", "replace": "変更後の文字列"}, ...]

書式の扱い:
  - find が単一 run に収まる段落 … その run を直接置換(部分書式も保たれる)。
  - find が複数 run にまたがる段落 … 段落先頭 run の書式で1つにまとめて置換する
    (段落内の部分書式は失われるが、段落スタイル=フォント等は保たれる)。
置換は **入力ファイルを変更せず**、出力ファイルに保存する。原本は不変。

使い方:
  python scripts/docx_replace.py 原本.docx --replacements changes.json -o 改訂版.docx
  python scripts/docx_replace.py 原本.docx -r changes.json          # 既定 原本-revised.docx

依存: python-docx。
"""
import argparse
import json
import os
import sys

for _s in (sys.stdout, sys.stderr):
    try:
        _s.reconfigure(encoding="utf-8")
    except (AttributeError, ValueError):
        pass


def load_replacements(path):
    try:
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)
    except OSError as e:
        raise SystemExit(f"error: --replacements を開けません: {e}")
    except ValueError as e:
        raise SystemExit(f"error: --replacements が不正な JSON です: {e}")
    if not isinstance(data, list):
        raise SystemExit('error: --replacements は [{"find": ..., "replace": ...}, ...] の JSON 配列にしてください')
    out = []
    for i, m in enumerate(data):
        if not isinstance(m, dict) or "find" not in m:
            raise SystemExit(f"error: --replacements[{i}] は find を持つオブジェクトにしてください")
        find = str(m["find"])
        if not find:
            raise SystemExit(f"error: --replacements[{i}] の find が空です")
        out.append((find, str(m.get("replace", ""))))
    return out


def all_paragraphs(doc):
    """本文 + 表セル(入れ子表も含む)の全段落を集める。"""
    paras = []

    def from_cell(cell):
        paras.extend(cell.paragraphs)
        for tbl in cell.tables:
            for row in tbl.rows:
                for c in row.cells:
                    from_cell(c)

    paras.extend(doc.paragraphs)
    for tbl in doc.tables:
        for row in tbl.rows:
            for c in row.cells:
                from_cell(c)
    return paras


def replace_in_paragraph(para, find, replace):
    """段落内の find を replace に置換し、置換した出現回数を返す。"""
    full = para.text
    if find not in full:
        return 0
    occurrences = full.count(find)
    runs = para.runs
    if len(runs) == 1:
        runs[0].text = full.replace(find, replace)  # 単一 run: 書式を保って置換
    elif runs:
        runs[0].text = full.replace(find, replace)  # 複数 run: 先頭 run に集約(段落書式は維持)
        for r in runs[1:]:
            r.text = ""
    else:  # run の無い段落(まれ)
        para.add_run(full.replace(find, replace))
    return occurrences


def main():
    p = argparse.ArgumentParser(description="原本 docx をコピーして合意済みの変更だけ差し替える")
    p.add_argument("docx", help="原本 docx(変更しない)")
    p.add_argument("-r", "--replacements", required=True, help='変更 JSON([{"find":...,"replace":...}])')
    p.add_argument("-o", "--output", help="出力 docx(省略時は <原本>-revised.docx)")
    a = p.parse_args()
    if not os.path.exists(a.docx):
        raise SystemExit(f"error: ファイルが見つかりません: {a.docx}")
    try:
        from docx import Document
    except ImportError:
        raise SystemExit("error: docx_replace には python-docx が必要です。`pip install python-docx` を実行してください")

    repls = load_replacements(a.replacements)
    doc = Document(a.docx)
    paras = all_paragraphs(doc)
    summary = []
    for find, replace in repls:
        total = sum(replace_in_paragraph(para, find, replace) for para in paras)
        summary.append((find, total))

    out = a.output or (os.path.splitext(a.docx)[0] + "-revised.docx")
    doc.save(out)
    n_zero = 0
    for find, total in summary:
        if total:
            sys.stderr.write(f"[OK]   {total}件: {find[:48]}\n")
        else:
            n_zero += 1
            sys.stderr.write(f"[WARN] 0件(原文に未検出): {find[:48]}\n")
    sys.stderr.write(f"wrote: {out}  (変更 {len(repls)} 種 / 未検出 {n_zero} 種、原本は不変)\n")


if __name__ == "__main__":
    main()
