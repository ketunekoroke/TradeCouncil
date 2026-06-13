#!/usr/bin/env python3
"""md_to_docx.py — Markdown を体裁付き docx に変換する CLI。

資料レビュー(document-review)で、改訂本文(Markdown)を「元と同形式(docx)」の成果物へ
書き出すためのヘルパー。見出し(#/##/###)、段落(インライン **太字**)、GFM 表、引用(>)、
水平線(---)を解釈する。最初の `#` は表題(中央・大きめ)、以降の `#` は見出し1。
日本語は明朝(本文)/ ゴシック(見出し)を eastAsia 指定で当てる。

⚠️ 注意(→ C2): この方式は Markdown からの**全面再構築**であり、元 docx の画像・チャート・
複雑な書式は保持されない。画像や体裁を保ちたい場合は scripts/docx_replace.py で原本を
コピー編集する方を使う(全面再構築は最終手段)。

使い方:
  python scripts/md_to_docx.py 改訂版.md                  # 改訂版.docx を生成
  python scripts/md_to_docx.py 改訂版.md -o out.docx

依存: python-docx。
環境変数: OFFICE_DOCX_BODY_FONT(既定 Yu Mincho)/ OFFICE_DOCX_HEAD_FONT(既定 Yu Gothic)。
          (旧 MAGI_DOCX_* も後方互換で読む — ADR-0011)
"""
import argparse
import os
import re
import sys

for _s in (sys.stdout, sys.stderr):
    try:
        _s.reconfigure(encoding="utf-8")
    except (AttributeError, ValueError):
        pass

JP_BODY = os.environ.get("OFFICE_DOCX_BODY_FONT") or os.environ.get("MAGI_DOCX_BODY_FONT") or "Yu Mincho"
JP_HEAD = os.environ.get("OFFICE_DOCX_HEAD_FONT") or os.environ.get("MAGI_DOCX_HEAD_FONT") or "Yu Gothic"


def _set_font(run, name):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:eastAsia", "w:ascii", "w:hAnsi"):
        rfonts.set(qn(attr), name)


def _add_inline(paragraph, text, font, base_bold=False):
    """**太字** を解釈しつつ runs を追加。"""
    bold = base_bold
    for seg in re.split(r"(\*\*)", text):
        if seg == "**":
            bold = not bold
            continue
        if seg == "":
            continue
        run = paragraph.add_run(seg)
        run.bold = bold
        _set_font(run, font)


def _is_table_sep(cells):
    return bool(cells) and all(re.fullmatch(r":?-{2,}:?", (c.strip() or "-")) for c in cells)


def convert(md_text, out_path):
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.size = Pt(10.5)
    rpr = normal.element.get_or_add_rPr()
    rfonts = OxmlElement("w:rFonts")
    for attr in ("w:eastAsia", "w:ascii", "w:hAnsi"):
        rfonts.set(qn(attr), JP_BODY)
    rpr.append(rfonts)

    lines = md_text.split("\n")
    i, n = 0, len(lines)
    title_used = False
    while i < n:
        st = lines[i].strip()
        # table block
        if st.startswith("|"):
            block = []
            while i < n and lines[i].strip().startswith("|"):
                block.append(lines[i].strip())
                i += 1
            rows = []
            for r in block:
                cells = [c.strip() for c in r.strip().strip("|").split("|")]
                if _is_table_sep(cells):
                    continue
                rows.append(cells)
            if rows:
                ncol = max(len(r) for r in rows)
                t = doc.add_table(rows=0, cols=ncol)
                t.style = "Table Grid"
                for ri, r in enumerate(rows):
                    cdoc = t.add_row().cells
                    for ci in range(ncol):
                        para = cdoc[ci].paragraphs[0]
                        _add_inline(para, r[ci] if ci < len(r) else "", JP_HEAD, base_bold=(ri == 0))
                        for run in para.runs:
                            run.font.size = Pt(9)
                doc.add_paragraph()
            continue
        m = re.match(r"^(#{1,3})\s+(.*)$", st)
        if m:
            level, txt = len(m.group(1)), m.group(2)
            if level == 1 and not title_used:
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run(txt)
                run.bold = True
                run.font.size = Pt(15)
                _set_font(run, JP_HEAD)
                title_used = True
            else:
                h = doc.add_heading(level=level)
                run = h.add_run(txt)
                _set_font(run, JP_HEAD)
            i += 1
            continue
        if re.fullmatch(r"-{3,}", st):
            i += 1
            continue
        if st.startswith(">"):
            qt = re.sub(r"^>\s?", "", st)
            if qt.strip():
                para = doc.add_paragraph()
                para.paragraph_format.left_indent = Pt(18)
                _add_inline(para, qt, JP_BODY)
                for run in para.runs:
                    run.font.size = Pt(9.5)
            i += 1
            continue
        if st == "":
            i += 1
            continue
        para = doc.add_paragraph()
        _add_inline(para, st, JP_BODY)
        i += 1
    doc.save(out_path)


def main():
    p = argparse.ArgumentParser(description="Markdown を体裁付き docx に変換")
    p.add_argument("path", help="入力 Markdown ファイル")
    p.add_argument("-o", "--output", help="出力 docx(省略時は入力と同名 .docx)")
    a = p.parse_args()
    if not os.path.exists(a.path):
        raise SystemExit(f"error: ファイルが見つかりません: {a.path}")
    try:
        import docx  # noqa: F401
    except ImportError:
        raise SystemExit("error: md_to_docx には python-docx が必要です。`pip install python-docx` を実行してください")
    out = a.output or (os.path.splitext(a.path)[0] + ".docx")
    with open(a.path, encoding="utf-8") as f:
        md = f.read()
    convert(md, out)
    sys.stderr.write(f"wrote: {out}\n")


if __name__ == "__main__":
    main()
